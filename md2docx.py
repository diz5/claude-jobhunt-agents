#!/usr/bin/env python3
"""md2docx.py — convert a Markdown résumé to .docx, cloning the look of a reference .docx.

Why: you keep your résumé as version-controllable Markdown, but applications want .docx.
This script takes a .docx you already like (your current résumé) as a STYLE REFERENCE and
re-emits your Markdown content in that exact look — fonts, sizes, bullet numbering, the
right-aligned date tab on job headers, and the education table if the reference has one.

Usage:
    python3 md2docx.py --ref "MyResume-old.docx" --in "resume.md" --out "resume.docx"

Markdown contract (see application-kit.example.md / examples for the shapes):
    # Name                              -> name line (style of the reference's 1st paragraph)
    contact line under the name         -> contact line (2nd paragraph's style)
    ## Section                          -> section heading (reference's Heading style)
    ### Role - Company | Date           -> job header (bold + right-aligned tab date)
    - bullet  /  "  - sub-bullet"       -> experience bullets (reference bullet numbering)
    - bullets under "## Skills"         -> skills bullets (reference skills-list style)
    **bold** spans                      -> bold runs
    Education lines: "**Degree** — School | Date" -> a table row if the reference has an
    education table, else plain lines.

How the reference is read (auto-detected, no configuration):
    name    = 1st non-empty paragraph          contact = 2nd non-empty paragraph
    heading = 1st Heading-style paragraph      body    = 1st plain paragraph after it
    job     = 1st bold paragraph containing a tab
    bullet  = 1st numbered paragraph not in List Paragraph style
    skills  = 1st numbered List Paragraph      table   = the reference's 1st table (education)
Anything not found falls back to the body style, so an unusual reference still converts.

Requires: python-docx  (pip install python-docx)
"""
import argparse
import copy
import re

from docx import Document
from lxml import etree

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


# ── reference harvesting ─────────────────────────────────────────────────────

def first_run_rpr(p_el):
    runs = p_el.findall(W + 'r')
    if runs:
        rpr = runs[0].find(W + 'rPr')
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def has_numpr(p):
    ppr = p._element.find(W + 'pPr')
    return ppr is not None and ppr.find(W + 'numPr') is not None


def first_run_bold(p):
    return bool(p.runs) and bool(p.runs[0].bold)


def harvest(doc):
    """Auto-detect prototype paragraphs in the reference document."""
    ps = [p for p in doc.paragraphs]
    nonempty = [p for p in ps if p.text.strip()]
    protos = {}
    protos['name'] = nonempty[0] if nonempty else None
    protos['contact'] = nonempty[1] if len(nonempty) > 1 else protos['name']
    protos['h1'] = next((p for p in ps if p.style.name.startswith('Heading')), protos['name'])
    heading_seen = False
    body = None
    for p in ps:
        if p.style.name.startswith('Heading'):
            heading_seen = True
        elif heading_seen and p.text.strip() and not has_numpr(p) and '\t' not in p.text:
            body = p
            break
    protos['body'] = body or protos['contact']
    protos['job'] = next((p for p in ps if '\t' in p.text and first_run_bold(p)), protos['body'])
    protos['bullet'] = next((p for p in ps if has_numpr(p) and p.style.name != 'List Paragraph'),
                            protos['body'])
    protos['skill'] = next((p for p in ps if has_numpr(p) and p.style.name == 'List Paragraph'),
                           protos['bullet'])
    elements = {k: copy.deepcopy(p._element) if p is not None else None for k, p in protos.items()}
    rprs = {k: (first_run_rpr(p._element) if p is not None else None) for k, p in protos.items()}
    table = copy.deepcopy(doc.tables[0]._element) if doc.tables else None
    return elements, rprs, table


# ── docx building blocks ─────────────────────────────────────────────────────

def blank_para(proto):
    el = copy.deepcopy(proto)
    for child in list(el):
        if child.tag != W + 'pPr':
            el.remove(child)
    return el


def set_bold(rpr, bold):
    for tag in ('b', 'bCs'):
        for e in rpr.findall(W + tag):
            rpr.remove(e)
    if bold:
        for tag in ('b', 'bCs'):
            etree.SubElement(rpr, W + tag)


def add_run(p, text=None, rpr=None, bold=None, tab=False, br=False):
    r = etree.SubElement(p, W + 'r')
    if rpr is not None:
        rp = copy.deepcopy(rpr)
        if bold is not None:
            set_bold(rp, bold)
        r.append(rp)
    if br:
        etree.SubElement(r, W + 'br')
    if tab:
        etree.SubElement(r, W + 'tab')
    if text:
        t = etree.SubElement(r, W + 't')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
    return r


def add_inline(p, md, rpr):
    """Emit runs for a markdown string, honoring **bold** spans."""
    for i, part in enumerate(re.split(r'\*\*', md)):
        if part:
            add_run(p, part, rpr, bold=(i % 2 == 1) or None)


def indent_more(p_el, delta=440):
    ppr = p_el.find(W + 'pPr')
    ind = ppr.find(W + 'ind') if ppr is not None else None
    if ind is not None and ind.get(W + 'left'):
        ind.set(W + 'left', str(int(ind.get(W + 'left')) + delta))


def clean(s):
    return re.sub(r'  +', ' ', s.replace('&nbsp;', ' ')).strip()


# ── education table ──────────────────────────────────────────────────────────

EDU_LINE = re.compile(r'^\*\*(?P<degree>.+?)\*\*\s*[—–-]\s*(?P<school>[^|]+?)(?:\s*\|\s*(?P<date>.+))?$')


def parse_edu(lines):
    rows = []
    for line in lines:
        m = EDU_LINE.match(clean(line))
        if m:
            rows.append((m['degree'].strip(), m['school'].strip(), (m['date'] or '').strip()))
    return rows


def build_table(tbl_proto, rows):
    tbl = copy.deepcopy(tbl_proto)
    trs = tbl.findall(W + 'tr')
    while len(trs) < len(rows):
        tbl.append(copy.deepcopy(trs[0]))
        trs = tbl.findall(W + 'tr')
    for extra in trs[len(rows):]:
        tbl.remove(extra)
    for row_el, (degree, school, date) in zip(tbl.findall(W + 'tr'), rows):
        for ci, tc in enumerate(row_el.findall(W + 'tc')):
            paras = tc.findall(W + 'p')
            keep = paras[0]
            rpr = first_run_rpr(keep)
            for extra in paras[1:]:
                tc.remove(extra)
            for child in list(keep):
                if child.tag != W + 'pPr':
                    keep.remove(child)
            if ci == 0:
                add_run(keep, degree, rpr, bold=True)
                if school:
                    add_run(keep, ' — ' + school, rpr, bold=False)
            else:
                add_run(keep, date, rpr, bold=True)
    return tbl


# ── conversion ───────────────────────────────────────────────────────────────

def convert(ref, md_path, out_path):
    doc = Document(ref)
    protos, rprs, tbl_proto = harvest(doc)
    body = doc.element.body
    for child in list(body):
        if child.tag in (W + 'p', W + 'tbl'):
            body.remove(child)
    sectpr = body.find(W + 'sectPr')

    def emit(el):
        if sectpr is not None:
            sectpr.addprevious(el)
        else:
            body.append(el)

    lines = open(md_path, encoding='utf-8').read().splitlines()
    section = None
    seen_name = seen_contact = False
    edu_buf = []

    def flush_edu():
        rows = parse_edu(edu_buf)
        if rows and tbl_proto is not None:
            emit(build_table(tbl_proto, rows))
        else:
            for line in edu_buf:
                p = blank_para(protos['body'])
                add_inline(p, clean(line), rprs['body'])
                emit(p)
        edu_buf.clear()

    for raw in lines:
        line = raw.rstrip()
        if not line.strip() or line.strip() == '---' or line.lstrip().startswith('<!--'):
            continue
        if line.startswith('# ') and not seen_name:
            p = blank_para(protos['name'])
            add_run(p, clean(line[2:]), rprs['name'])
            emit(p)
            seen_name = True
        elif seen_name and not seen_contact and not line.startswith('#'):
            p = blank_para(protos['contact'])
            add_run(p, clean(line), rprs['contact'])
            emit(p)
            seen_contact = True
        elif line.startswith('## '):
            if section == 'education':
                flush_edu()
            title = clean(line[3:])
            section = 'education' if title.lower() == 'education' else (
                'skills' if title.lower() == 'skills' else 'normal')
            p = blank_para(protos['h1'])
            add_run(p, title, rprs['h1'])
            emit(p)
        elif section == 'education':
            edu_buf.append(line)
        elif line.startswith('### '):
            head = clean(line[4:]).replace(' — ', ' – ')
            role, _, date = head.partition(' | ')
            p = blank_para(protos['job'])
            add_run(p, role.strip(), rprs['job'])
            add_run(p, None, rprs['job'], tab=True)
            add_run(p, date.strip(), rprs['job'])
            emit(p)
        elif re.match(r'^\s{2,}- ', line):
            p = blank_para(protos['bullet'])
            indent_more(p)
            add_inline(p, clean(re.sub(r'^\s+- ', '', line)), rprs['bullet'])
            emit(p)
        elif line.startswith('- '):
            proto = 'skill' if section == 'skills' else 'bullet'
            p = blank_para(protos[proto])
            add_inline(p, clean(line[2:]), rprs[proto])
            emit(p)
        else:
            p = blank_para(protos['body'])
            add_inline(p, clean(line), rprs['body'])
            emit(p)
    if section == 'education':
        flush_edu()
    doc.save(out_path)
    print(f"wrote {out_path}")


def main():
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument('--ref', required=True, help='reference .docx whose look to clone')
    ap.add_argument('--in', dest='src', required=True, help='input resume .md')
    ap.add_argument('--out', required=True, help='output .docx path')
    args = ap.parse_args()
    convert(args.ref, args.src, args.out)


if __name__ == '__main__':
    main()
